"""
Genera un document Word amb els 4 articles llargs del pla de comunicació Criteri ESG (agost-setembre 2026):
1. "El mapa dels 16 estàndards ESG" (12 agost, ~700 paraules)
2. "Per què els directors de sostenibilitat perden el 60% del seu temps" (9 setembre, ~1.500 paraules)
3. "Més enllà del Checkbox #2" (15 setembre, ~500 paraules)
4. "El semàfor metodològic: com avaluem el que llegim" (23 setembre, ~800 paraules)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Paleta Criteri ESG
COLOR_PRIMARY = RGBColor(0x2C, 0x18, 0x10)    # marró molt fosc
COLOR_SECONDARY = RGBColor(0x5C, 0x3A, 0x1E)  # marró fosc
COLOR_ACCENT = RGBColor(0xB8, 0x73, 0x33)     # coure
COLOR_ACCENT_LIGHT = RGBColor(0xE8, 0xC9, 0x9A)  # coure clar
COLOR_CREAM = RGBColor(0xF5, 0xEF, 0xE6)      # crema
COLOR_MUTED = RGBColor(0x8B, 0x73, 0x55)      # text secundari
COLOR_HOVER = RGBColor(0x8A, 0x55, 0x26)      # accent fosc

# Fonts (aproximacions a Fraunces/Inter/JetBrains Mono disponibles al sistema)
FONT_SERIF = "Georgia"        # proper a Fraunces
FONT_SANS = "Calibri"         # proper a Inter
FONT_MONO = "Consolas"        # proper a JetBrains Mono

doc = Document()

# Pàgina A4 amb marges raonables
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== Estils base =====
styles = doc.styles
normal = styles['Normal']
normal.font.name = FONT_SANS
normal.font.size = Pt(11)
normal.font.color.rgb = COLOR_PRIMARY
normal.paragraph_format.line_spacing = 1.4
normal.paragraph_format.space_after = Pt(8)

# Funció helper per afegir paràgraf amb format
def add_paragraph(text, font=FONT_SANS, size=11, color=COLOR_PRIMARY, bold=False, italic=False,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=8, line_spacing=1.4,
                   first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

def add_eyebrow(text):
    """Petita etiqueta en mono a dalt d'una secció"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.name = FONT_MONO
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_HOVER
    run.bold = True
    # Letter spacing
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), '40')  # 40 twips ≈ 2pt
    rPr.append(spacing)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT_SERIF
    run.font.size = Pt(24)
    run.font.color.rgb = COLOR_PRIMARY
    run.bold = True
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT_SERIF
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_PRIMARY
    run.bold = True
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_HOVER
    run.bold = True
    return p

def add_body(text, italic=False):
    """Cos de text normal amb indentació de primera línia"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # NO justificat segons regla del CONTEXT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_PRIMARY
    run.italic = italic
    return p

def add_body_no_indent(text, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.45
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_PRIMARY
    run.italic = italic
    return p

def add_blockquote(text):
    """Cita destacada amb barra esquerra coure"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.4
    # Afegir barra esquerra via border
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')  # 3pt
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'B87333')  # coure
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = FONT_SERIF
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_SECONDARY
    run.italic = True
    return p

def add_meta(text):
    """Text meta en mono (per dates, fonts, etc.)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = FONT_MONO
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED
    return p

def add_question_block(text):
    """Pregunta final destacada"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.line_spacing = 1.4
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT_SERIF
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_HOVER
    run.italic = True
    return p

def add_separator():
    """Línia separadora coure"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # 0.75pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B87333')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

# ===== PORTADA =====
# Logotip textual
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(120)
p.paragraph_format.space_after = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Criteri ESG")
run.font.name = FONT_SERIF
run.font.size = Pt(36)
run.font.color.rgb = COLOR_PRIMARY
run.bold = True
# Punt coure
run2 = p.add_run(".")
run2.font.name = FONT_SERIF
run2.font.size = Pt(36)
run2.font.color.rgb = COLOR_ACCENT
run2.bold = True

# Subtítol
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(80)
run = p.add_run("Inteligencia ESG para decisiones éticas")
run.font.name = FONT_SANS
run.font.size = Pt(13)
run.font.color.rgb = COLOR_MUTED
run.italic = True

# Títol del document
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Articles llargs")
run.font.name = FONT_SERIF
run.font.size = Pt(28)
run.font.color.rgb = COLOR_PRIMARY
run.bold = True

# Subtítol
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(40)
run = p.add_run("Pla de comunicació · agost-setembre 2026")
run.font.name = FONT_SERIF
run.font.size = Pt(16)
run.font.color.rgb = COLOR_HOVER
run.italic = True

# Meta
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("4 articles · 3.500 paraules aproximadament")
run.font.name = FONT_MONO
run.font.size = Pt(10)
run.font.color.rgb = COLOR_MUTED

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Veus: Criteri ESG (3 articles) · Paolo G. (article 2, veu personal)")
run.font.name = FONT_MONO
run.font.size = Pt(10)
run.font.color.rgb = COLOR_MUTED

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Idioma: castellà per publicació · català per a guions interns")
run.font.name = FONT_MONO
run.font.size = Pt(10)
run.font.color.rgb = COLOR_MUTED

# Línia accent coure
add_separator()

# Taula d'índex
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
run = p.add_run("Índex")
run.font.name = FONT_SERIF
run.font.size = Pt(16)
run.font.color.rgb = COLOR_PRIMARY
run.bold = True

index_entries = [
    ("Article 1", "El mapa de los 16 estándares ESG", "12 agost 2026 · LinkedIn + Web", "~700 paraules"),
    ("Article 2", "Por qué los directores de sostenibilidad pierden el 60% de su tiempo", "9 setembre 2026 · LinkedIn + Article web", "~1.500 paraules"),
    ("Article 3", "Más allá del Checkbox #2", "15 setembre 2026 · LinkedIn + Web", "~500 paraules"),
    ("Article 4", "El semáforo metodológico: cómo evaluamos lo que leemos", "23 setembre 2026 · LinkedIn + Article web", "~800 paraules"),
]
for num, title, meta, length in index_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{num}. ")
    run.font.name = FONT_MONO
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_ACCENT
    run.bold = True
    run = p.add_run(title)
    run.font.name = FONT_SERIF
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_PRIMARY
    run.bold = True
    # Meta en línia nova
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(1.2)
    run = p.add_run(f"{meta} · {length}")
    run.font.name = FONT_MONO
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED

# Salta pàgina
add_page_break()

# ========================================
# ARTICLE 1 — El mapa dels 16 estàndards ESG
# ========================================
add_eyebrow("Article 1 · 12 agost 2026 · LinkedIn + Web")
add_h1("El mapa de los 16 estándares ESG que ningún director debería confundir")

add_meta("Por Criteri ESG · 6 minuts de lectura · Font: criteriesg.com/estandares-esg")

add_body_no_indent("Hay una confusión que aparece en casi todas las conversaciones que mantenemos con directores de sostenibilidad. No es sobre el CSRD, ni sobre el CSDDD, ni sobre EcoVadis. Es sobre qué es cada uno de esos tres nombres. Y esa confusión tiene consecuencias operativas reales: decisiones que se toman con el marco equivocado, presupuestos que se asignan a la certificación incorrecta, equipos que dedican semanas a un documento que no les corresponde.")

add_body("Esta semana publicamos el mapa de los 16 estándares ESG con los que trabaja un director de sostenibilidad español. Lo hemos organizado en tres categorías, con tres colores distintos, porque la distinción entre ellas importa más que la lista en sí.")

add_h2("Tres categorías, tres colores")

add_body_no_indent("No todos los estándares son lo mismo. Algunos te obligan por ley. Otros te orientan sobre cómo reportar. Otros te evalúan desde fuera. Confundir un tipo con otro tiene costes: el más común es invertir recursos en una certificación cuando lo que tu empresa necesita es cumplir una regulación.")

add_h3("Regulaciones — las que te obligan")
add_body("Cinco elementos: CSRD/ESRS, CSDDD, SFDR, Taxonomía UE, EMAS. Son leyes europeas o estatales. No las eliges: las recibes. El incumplimiento tiene sanción administrativa. Su carácter vinculante las convierte en el suelo mínimo sobre el que todo lo demás se construye. En el mapa, las identificamos con un marrón muy oscuro (#5C3A1E) porque son la base: lo que está debajo de todo.")

add_h3("Frameworks — los que te orientan")
add_body("Cinco elementos: GRI, SASB, TNFD, TCFD, ISO 26000. Son estándares de reporting. Te dicen cómo contar lo que haces, no qué debes hacer. No son obligatorios, pero la Comisión Europea los referencia en los ESRS: usarlos bien te ahorra trabajo y te alinea con lo que vendrá. En el mapa, coure (#B87333): un color más cálido, porque son herramientas de trabajo, no obligaciones.")

add_h3("Certificaciones y ratings — los que te evalúan")
add_body("Seis elementos: EcoVadis, B Corp, MSCI ESG, CDP, SGE 21, Sustainalytics. Son evaluaciones externas que clientes, inversores y socios utilizan para posicionarte. Tienen lógica de mercado: las pides porque te abren puertas, no porque te obliguen. En el mapa, coure clar (#E8C99A): un color más suave, porque su valor depende de cómo los uses, no de que los tengas.")

add_blockquote("Una regulación te obliga. Un framework te orienta. Una certificación te evalúa. Confundirlas tiene consecuencias operativas. No las confundáis.")

add_h2("Por qué esta clasificación importa")

add_body_no_indent("El error más frecuente que hemos visto en consultorías y direcciones de sostenibilidad es tratar los 16 estándares como una lista plana. \"Tenemos que estar en todos\". No. No tenéis que estar en todos. Tenéis que saber en cuáles debéis estar por obligación, en cuáles os conviene estar por coherencia, y en cuáles os evaluarán quieran o no.")

add_body("Un ejemplo concreto: una empresa industrial española de 400 empleados, con EcoVadis Plata y B Corp, recibe la visita de un cliente alemán que le pide TCFD. La reacción instintiva es \"tenemos que certificarnos en TCFD\". Pero TCFD no es una certificación: es un framework. No se certifica, se reporta. Y si ya reportas con GRI, gran parte del trabajo está hecho. La diferencia entre \"certificarse\" y \"reportar\" puede ser 6 meses de trabajo y 30.000€ de coste innecesario.")

add_body("Otro ejemplo: una empresa cotizada que recibe presión de inversores por su rating MSCI. La reacción instintiva es \"tenemos que subir MSCI\". Pero MSCI es un rating basado en información pública. Subirlo no requiere certificar nada nuevo: requiere comunicar mejor lo que ya haces. La acción correcta no es contratar a MSCI: es revisar tu discurso y asegurarte de que lo que publicas en tu memoria de sostenibilidad (probablemente GRI) está alineado con lo que MSCI busca.")

add_h2("Cómo cruzarlos")

add_body_no_indent("El mapa no es solo una clasificación. Es una herramienta de cross-reference. Cada informe que publicamos en Criteri ESG indica qué estándares se ven afectados y con qué intensidad. Si la Comisión Europea revisa los ESRS, eso afecta al reporting CSRD pero también al framework GRI (por convergencia) y al rating EcoVadis (que premia el reporting GRI). Si MSCI baja el rating del sector utilities español, eso puede estar conectado con datos TCFD y con la Taxonomía UE.")

add_body("El cross-reference es la parte menos visible y más útil de nuestro trabajo. Te ahorra tener que leer tú mismo las 47 páginas del informe europeo para saber si tu certificación EcoVadis se ve afectada. Te lo decimos nosotros, con la fuente.")

add_h2("Un mapa vivo")

add_body_no_indent("Este mapa no es definitivo. Lo actualizaremos cada trimestre, porque el paisaje ESG cambia: el BCE publicará nuevas guías, B Lab modificará el B Impact Assessment, la Comisión Europea revisará los ESRS otra vez. Cuando eso pase, actualizaremos el mapa y lo comunicaremos en la newsletter.")

add_body("Si trabajas en sostenibilidad y quieres usar este mapa internamente, descárgalo. Es gratis. La única condición: si lo compartes, cita la fuente. No porque nos importe el crédito, sino porque la fuente es lo que distingue un mapa fiable de uno improvisado.")

add_question_block("Pregunta para mejorar: ¿sabrías decir, sin mirar el mapa, en qué categoría está cada uno de los 16 estándares que afectan a tu empresa?")

add_separator()

# Salta pàgina al següent article
add_page_break()

# ========================================
# ARTICLE 2 — Per què els directors perden el 60% del seu temps
# ========================================
add_eyebrow("Article 2 · 9 setembre 2026 · Veu personal de Paolo G.")
add_h1("Por qué los directores de sostenibilidad pierden el 60% de su tiempo recopilando información")

add_meta("Por Paolo G., fundador de Criteri ESG · 12 minuts de lectura · Dades: estudi intern basat en 12 entrevistes amb directors de sostenibilitat espanyols")

add_body_no_indent("Hace ocho meses dejé mi trabajo para construir Criteri ESG. Lo hice por una razón concreta: durante años fui el director de sostenibilidad que recibía 14 informes institucionales por semana y no tenía tiempo de leerlos. Y porque la solución no era más información: era criterio.")

add_body("En los últimos meses he hablado con 12 directores de sostenibilidad de empresas medianas y grandes en España. La conversación siempre empieza igual: \"¿cómo os organizáis para procesar todo lo que llega?\". La respuesta, también siempre igual: \"mal\".")

add_h2("La distribución del tiempo")

add_body_no_indent("Antes de pedirles que hablaran, les pedí que estimaran cuánto tiempo dedicaban a cuatro tareas: recopilar información, leer informes, analizar y decidir. Los resultados, agregados y redondeados:")

add_blockquote("Recopilar información: 60%. Leer informes: 15%. Analizar: 15%. Decidir: 10%.")

add_body("El dato es estimativo, no científico. Pero es consistente. Doce personas, en sectores distintos (banca, industria, energía, consultoría, distribución), con tamaños de empresa distintos, describen un patrón similar: la mayoría de su tiempo se va en buscar, descargar, clasificar y archivar información externa. No en pensar qué hacer con ella.")

add_body("La paradoja del reporting ESG es que más información no genera más decisiones. Genera más ansiedad y menos capacidad de priorizar. El director que recibe 14 informes semanales no toma 14 decisiones: toma una, tarde, con un solo informe, basándose en una memoria imperfecta de los otros 13.")

add_h2("Tres casos anonimizados")

add_h3("Caso 1 · Directora de sostenibilidad, empresa industrial, 600 empleados")
add_body("\"El lunes llega el ESRS Q&A Platform con 8 nuevas respuestas. El martes, un informe de MSCI sobre el sector. El miércoles, una actualización de EcoVadis. El jueves, un press release de Sustainalytics. El viernes, algo del BCE. Yo paso el lunes entero descargando. El martes leyendo titulares. El miércoles contestando emails. El jueves preparando la reunión del comité. Y el viernes me doy cuenta de que no he pensado nada en toda la semana.\"")

add_h3("Caso 2 · Director de RSC, entidad financiera, 2.000 empleados")
add_body("\"Tenemos una suscripción a Sustainalytics, otra a MSCI, otra a Refinitiv. Cada una con su dashboard. Y luego están las fuentes públicas: ESRS, BCE, EBA, ESMA. Mi equipo dedica a una persona a tiempo completo solo a monitorizar. La semana pasada me trajo un resumen de 12 páginas. Yo lo leí en el metro. ¿Decisión tomada? Ninguna. Solo información que entra y sale.\"")

add_h3("Caso 3 · Responsable de reporting, consultoría ESG, 80 empleados")
add_body("\"Para cada cliente tengo que leer el mismo informe de tres maneras distintas. Una para el cliente A que tiene EcoVadis Plata. Otra para el cliente B que tiene B Corp. Otra para el cliente C que solo hace CSRD. El informe es el mismo. El análisis que hago yo, en cambio, es tres veces. Si alguien me diera un cross-reference ya hecho, ahorraría 12 horas a la semana.\"")

add_h2("El problema no es la información. Es la falta de criterio.")

add_body_no_indent("Cuando le digo a un director \"tienes 14 informes en la bandeja\", la reacción instintiva es \"necesito más tiempo\". Pero más tiempo no resuelve nada: el lunes siguiente llegarán 14 más. Lo que falta no es tiempo. Es criterio: saber cuáles de los 14 importan para tu empresa, qué dicen realmente y qué tienes que hacer con ellos.")

add_body("El criterio no es una virtud personal. Es una estructura. Es saber que un informe de 47 páginas del BCE se reduce a tres ideas operativas. Es saber que la revisión de los ESRS te afecta de manera distinta si tienes EcoVadis Plata que si tienes B Corp. Es saber que el CSDDD cambia tu matriz de riesgo de proveedores aunque tu empresa no esté en el umbral de los 1.000 empleados.")

add_body("El criterio se construye con tres pasos: filtrar (qué informes merecen tu atención), sintetizar (qué dicen realmente, en 7 minutos) y cruzar (con las certificaciones que ya tienes o que te faltan). Nada más. Si no haces los tres, no tienes criterio: tienes una bandeja de entrada llena.")

add_h2("Una propuesta concreta")

add_body_no_indent("No voy a proponer que leas más. Voy a proponer lo contrario: que leas menos, pero con criterio.")

add_body("Criteri ESG no es un agregador. No recopilamos todos los informes para que tú los leas. Eso ya lo tienes. Lo que hacemos es seleccionar los 4-6 informes más relevantes de cada quincena, sintetizarlos en 7 minutos de lectura, cruzarlos con tus certificaciones (EcoVadis, B Corp, MSCI, GRI, SGE 21, las que tengas) y recomendarte tres acciones concretas. Nada más.")

add_body("No es magia. Es criterio aplicado. Y se basa en algo simple: si un director de sostenibilidad dedicara el 60% de su tiempo a pensar en lugar de a recopilar, sus decisiones serían mejores. Y si las decisiones son mejores, el reporting mejora, los ratings mejoran, los comités de sostenibilidad dejan de ser trámites. El cambio no empieza por más datos: empieza por más tiempo para procesarlos.")

add_h2("Lo que no resolveremos")

add_body_no_indent("Criteri ESG no resuelve todos los problemas de un director de sostenibilidad. No sustituye la conversación con el consejo. No redacta la memoria de sostenibilidad (todavía). No decide por ti qué certificaciones priorizar. No negocia con EcoVadis cuando el rating baja. No convence al CFO de que el presupuesto de sostenibilidad no se toca.")

add_body("Lo que sí resolvemos: el 60% del tiempo que se va en recopilar, leer y estructurar. Las 5 horas semanales que un director dedica a tareas que no aportan valor estratégico. El tiempo que deberías estar dedicando a pensar.")

add_body("Si reconoces esta situación, te entendemos. Estamos aquí para devolverte el 60%.")

add_question_block("Pregunta para mejorar: ¿cuántas horas a la semana dedicas a leer informes que, una semana después, no recuerdas qué decían?")

add_separator()
add_page_break()

# ========================================
# ARTICLE 3 — Més enllà del Checkbox #2
# ========================================
add_eyebrow("Article 3 · 15 setembre 2026 · Carta ètica quinzenal #2")
add_h1("Más allá del Checkbox #2")

add_meta("Por Criteri ESG · 3 minuts de lectura · Cada quinzena, una pregunta")

add_body_no_indent("Pregunta de esta quincena:")

add_blockquote("Si mañana te pidieran demostrar el impacto positivo de tu empresa con un solo dato, ¿cuál elegirías y por qué?")

add_body_no_indent("La mayoría de directores de sostenibilidad no sabe responder. Tienen 47 KPIs y ninguno que sintetice. Tienen métricas de emisiones Scope 1, 2 y 3. Tienen porcentaje de mujeres en el comité de dirección. Tienen horas de formación en ética por empleado. Tienen tasa de accidentes laborales. Tienen cobertura de auditoría social. Tienen todo. Y nada.")

add_body("Tener demasiadas métricas sin un relato coherente es tan inútil como no tener ninguna. La cuestión no es cuantificar más, es saber qué cuenta. Y qué cuenta no es lo que más fácil es medir: es lo que más fielmente refleja el valor real de tu empresa para las personas que se ven afectadas por ella.")

add_h2("Dos ejemplos anonimizados")

add_h3("Empresa A · 1.200 empleados · sector industrial")
add_body("\"Nuestro KPI estrella es la reducción del 18% en emisiones Scope 1 y 2 en tres años. Lo ponemos en la portada de la memoria. Lo comunicamos en cada earnings call. Lo celebramos internamente. Pero cuando nos preguntan por el Scope 3 —que es donde está el 87% de nuestra huella—, cambiamos de tema.\"")

add_h3("Empresa B · 350 empleados · sector servicios")
add_body("\"No tenemos un KPI estrella. Tenemos 30 indicadores que reportamos con GRI. Pero si tuviera que elegir uno, elegiría la tasa de retención de talento en zonas rurales donde somos el principal empleador. Es el dato que mejor explica por qué la comunidad nos defiende cuando alguien ataca el proyecto. Y no lo reportamos en la memoria, porque no encaja en ninguna categoría GRI estándar.\"")

add_h2("La diferencia entre los dos")

add_body_no_indent("La empresa A tiene un dato potente que no cuenta toda la verdad. La empresa B tiene un dato humilde que cuenta exactamente la verdad que importa. La empresa A comunica mejor. La empresa B decide mejor.")

add_body("La pregunta de esta quincena no es retórica. Es operativa. Si no sabes qué dato elegir, puedes estar sufriendo uno de estos dos problemas: o tienes los datos equivocados (los fáciles de medir en lugar de los importantes), o tienes el relato equivocado (no has pensado cuál es el valor real de tu empresa para las personas que la rodean).")

add_body("En ambos casos, la solución no es medir más. Es pensar mejor. Y para pensar mejor, a veces hace falta hacer la pregunta incómoda: ¿qué dato, si desapareciera mañana, dejaría un vacío real en alguien que no sea accionista?")

add_question_block("Y si no sabes qué dato elegir, ¿es el dato equivocado o el relato equivocado?")

add_separator()
add_page_break()

# ========================================
# ARTICLE 4 — El semàforo metodològic
# ========================================
add_eyebrow("Article 4 · 23 setembre 2026 · LinkedIn + Article web")
add_h1("El semáforo metodológico: cómo evaluamos lo que leemos")

add_meta("Por Criteri ESG · 7 minuts de lectura · Font: metodologia interna Criteri ESG")

add_body_no_indent("No todos los informes se crean igual. Un informe de 80 páginas del Banco Central Europeo sobre riesgo climático no es comparable con un press release de 3 páginas de MSCI anunciando un cambio metodológico. Y, sin embargo, los directores de sostenibilidad los reciben en la misma bandeja de entrada, los tratan con el mismo protocolo y les dedican el mismo tiempo. Esa igualdad de trato es un error.")

add_body("Por eso hemos creado un semáforo metodológico: una herramienta interna que aplicamos a cada informe que procesamos antes de publicarlo. No evalúa si el informe es bueno o malo. Evalúa si es útil para tomar decisiones operativas en una empresa concreta. Y lo hace en cinco dimensiones.")

add_h2("Cinco dimensiones")

add_h3("1. Scope 3")
add_body("¿El informe cubre emisiones indirectas de la cadena de valor o solo emisiones directas (Scope 1 y 2)? En la mayoría de sectores, el Scope 3 representa entre el 70% y el 90% de la huella total. Un informe que lo ignora es como un análisis de sangre que solo mira el brazo: técnicamente correcto, prácticamente inservible. Nota verde si el Scope 3 está integrado y desagregado por categoría. Nota amarilla si se menciona pero no se cuantifica. Nota roja si se ignora.")

add_h3("2. Plazos")
add_body("¿Las fechas que menciona el informe son operativas (con calendario concreto) o aspiracionales (\"para 2050\", \"en los próximos años\")? Las fechas aspiracionales no son malas: son políticas. Pero un director de sostenibilidad que tiene que planificar inversiones para 2027 no puede trabajar con \"para 2050\". Nota verde si hay calendario con años concretos y hitos intermedios. Nota amarilla si hay fechas finales pero no hitos. Nota roja si todo es aspiracional.")

add_h3("3. Fuentes")
add_body("¿Las afirmaciones del informe están referenciadas (con citas, enlaces, metodología accesible) o son afirmaciones sin respaldo? El greenwashing vive de las afirmaciones sin fuente. Un informe que dice \"las empresas europeas están mejorando su gestión de residuos\" sin citar qué empresas, qué periodo y qué datos es, metodológicamente, propaganda. Nota verde si cada afirmación lleva fuente verificable. Nota amarilla si hay fuentes pero no son accesibles o están desactualizadas. Nota roja si no hay fuentes.")

add_h3("4. Granularidad")
add_body("¿El informe ofrece datos desagregados por sector, geografía y tamaño de empresa, o solo agregados a nivel europeo? Un director de sostenibilidad de una PIME industrial en Catalunya no puede usar un dato agregado de la UE. Necesita el dato del sector, de la región, del rango de empleados. La granularidad es lo que convierte un informe institucional en una herramienta operativa. Nota verde si hay desagregación triple. Nota amarilla si hay desagregación por sector pero no por geografía o tamaño. Nota roja si todo es agregado.")

add_h3("5. Verificación")
add_body("¿El informe está auditado externamente o es auto-reportado? La auto-verificación no es necesariamente mala: muchas fuentes fiables se auto-reportan (MSCI, Sustainalytics). Pero un director de sostenibilidad debe saber si lo que lee ha pasado por un filtro independiente o no. Nota verde si hay auditoría externa de terceros con metodología pública. Nota amarilla si hay revisión interna pero no externa. Nota roja si es auto-reportado sin ningún tipo de verificación.")

add_h2("La nota final")

add_body_no_indent("Cada dimensión recibe una nota (verde, amarilla, roja). La nota final del informe es el resultado agregado, con una escala de cinco niveles:")

add_meta("A · Robusto fuerte · 5 verdes")
add_meta("B · Robusto · 4 verdes + 1 amarilla")
add_meta("C · Débil · combinación de verdes, amarillas y hasta 1 roja")
add_meta("D · Insuficiente · 2 o más rojas")

add_blockquote("Cada informe que publicamos lleva su semáforo. Es nuestra forma de no confundir rigor con volumen.")

add_h2("Un ejemplo concreto")

add_body_no_indent("Esta semana la Comisión Europea publicó la revisión de los ESRS. 47 páginas. Nuestra evaluación:")

add_meta("Scope 3: amarillo (se menciona pero se simplifica)")
add_meta("Plazos: amarillo (fechas operativas para reporting, pero sin hitos para transición)")
add_meta("Fuentes: verde (cada afirmación lleva referencia al ESRS original)")
add_meta("Granularidad: rojo (todo agregado a nivel UE, sin desagregar por sector o tamaño)")
add_meta("Verificación: amarillo (auditoría externa prevista pero no detallada)")

add_body("Nota final: C · Débil. El informe es riguroso en su arquitectura pero deja fuera el detalle granular que las empresas necesitan para implementar. Si tienes EcoVadis Plata, te afecta de una manera (deberás revisar tu matriz de materialidad). Si tienes B Corp, te afecta de otra (tu B Impact Assessment no cambia, pero tu reporting CSRD sí). Eso es lo que hacemos en Criteri ESG: cruzar el informe con tus certificaciones y decirte qué hacer.")

add_h2("Por qué publicamos el semáforo")

add_body_no_indent("Podríamos haber guardado esta metodología interna. No lo hacemos por tres razones. Primero, porque la transparencia es nuestra principal arma contra el greenwashing: si decimos que un informe es \"C · Débil\", cualquiera puede ver cómo llegamos a esa nota. Segundo, porque el semáforo ayuda a los directores a priorizar: si un informe es \"D · Insuficiente\", quizá no merezca tus 30 minutos. Tercero, porque nos obliga a ser consistentes: si aplicamos el semáforo a los demás, también lo aplicamos a nosotros. Cuando publiquemos un análisis Criteri ESG, llevará su propio semáforo.")

add_body("El semáforo no es un juicio moral. Es un juicio operativo. Un informe \"D · Insuficiente\" puede ser útil para entender el contexto político europeo; simplemente no es útil para decidir si debes invertir 200.000€ en un nuevo sistema de reporting. Distinguimos entre lo que merece ser leído y lo que merece ser actuado. Y eso, en un mundo de información infinita, ya es un acto de criterio.")

add_question_block("Pregunta para mejorar: ¿aplicas algún criterio metodológico a los informes que recibes, o los tratas a todos por igual?")

# ===== FOOTER =====
add_separator()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Criteri ESG · Inteligencia ESG para decisiones éticas · criteriesg.com")
run.font.name = FONT_MONO
run.font.size = Pt(9)
run.font.color.rgb = COLOR_MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Pla de comunicació agost-setembre 2026 · Articles llargs · v1.0")
run.font.name = FONT_MONO
run.font.size = Pt(8)
run.font.color.rgb = COLOR_MUTED

# Guarda
output_path = "/home/z/my-project/download/articles-llargs-criteri-esg-agost-setembre-2026.docx"
doc.save(output_path)
print(f"✓ Word guardat a {output_path}")

# Comptar paraules
import zipfile, re
with zipfile.ZipFile(output_path) as z:
    with z.open('word/document.xml') as f:
        content = f.read().decode('utf-8')
# Elimina tags XML per comptar paraules
text = re.sub(r'<[^>]+>', ' ', content)
words = len(text.split())
print(f"  Paraules totals (aprox): {words}")
